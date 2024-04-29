import os
from flask import Flask, render_template, request, redirect, url_for, send_file
from PyPDF2 import PdfReader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv
from flaskext.markdown import Markdown
from docx import Document
from io import BytesIO

load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

app = Flask(__name__)
Markdown(app)

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
PDF_FOLDER_PATH = os.path.join(SCRIPT_DIR, "pdf")


def get_pdf_text(pdf_doc):
    text = ""
    pdf_reader = PdfReader(pdf_doc)
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def get_text_chunks(text):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=10000, chunk_overlap=3000)
    chunks = text_splitter.split_text(text)
    return chunks

def get_vector_store(text_chunks, file_name):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    vector_store = FAISS.from_texts(text_chunks, embedding=embeddings)
    vector_store.save_local(f"faiss_index_{file_name}")

def get_conversational_chain():
    prompt_template = """ The given information is about resume of a person. Be specific and answer the question based on the information provided. If no information is found return "No information found". Be crisp and give answer.

    Context:
    {context}

    Question:
    {question}

    Answer:"""

    model = ChatGoogleGenerativeAI(model="gemini-pro", temperature=0.5)
    prompt = PromptTemplate(template=prompt_template, input_variables=["context", "question"])
    chain = load_qa_chain(model, chain_type="stuff", prompt=prompt)
    return chain

def user_input(user_question, file_name):
    embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")
    new_db = FAISS.load_local(f"faiss_index_{file_name}", embeddings, allow_dangerous_deserialization=True)
    if user_question is None:
        docs = []  # If no question is provided, use an empty list of docs
    else:
        docs = new_db.similarity_search(user_question)
    chain = get_conversational_chain()
    response = chain({"input_documents": docs, "question": user_question}, return_only_outputs=True)
    return response["output_text"]

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        user_question = request.form["question"]
        answers = []
        if not os.path.exists(PDF_FOLDER_PATH):
            os.makedirs(PDF_FOLDER_PATH)
        pdf_files = [os.path.join(PDF_FOLDER_PATH, f) for f in os.listdir(PDF_FOLDER_PATH) if f.endswith(".pdf")]
        for pdf_file in pdf_files:
            file_name = os.path.basename(pdf_file).split(".")[0]
            raw_text = get_pdf_text(pdf_file)
            text_chunks = get_text_chunks(raw_text)
            get_vector_store(text_chunks, file_name)
            answer = user_input(user_question, file_name)
            answers.append((file_name, answer))
        return render_template("index.html", question=user_question, answers=answers)
    return render_template("index.html")

@app.route("/process_pdfs", methods=["POST"])
def process_pdfs():
    if not os.path.exists(PDF_FOLDER_PATH):
        os.makedirs(PDF_FOLDER_PATH)
    return redirect(url_for("index"))

@app.route("/download_answers", methods=["POST"])
def download_answers():
    selected_answers = request.form.getlist("selected_answers")
    document = Document()

    for file_name in selected_answers:
        answer = request.form.get(f"answer-{file_name}")
        if answer:
            document.add_heading(f"{file_name}.pdf", level=0)
            document.add_paragraph(answer)

    file_stream = BytesIO()
    document.save(file_stream)
    file_stream.seek(0)

    return send_file(
        file_stream,
        as_attachment=True,
        download_name="selected_answers.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

if __name__ == "__main__":
    app.run(debug=True, port = 5001)